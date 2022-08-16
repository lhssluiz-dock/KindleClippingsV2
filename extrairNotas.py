from Handlers.FilesHandler import *
from Handlers.NotesHandler import *
from Services.cleanDictionary import *
from Services.cleanPageLine import *
from Services.getPage import *
from Services.getPosition import *
from Services.cleanPositionLine import *
from Handlers.DocumentHandler import *

import operator
from docx.shared import Inches

file_path = 'C:/Users/luiz.sa/Desktop/ExtrairNotasKindle/My Clippings.txt'
titleToSearch = input('Insira o nome do título do livro igual está no Kindle: ')
pageOrPosition = input('Clique 1: Para Posição \nClique 2: para Página \nSua Escolha:  ')
linePosition = 0
arrayLines = []
dictPositionNotes, lastDict = {}, {}

arrayLines = openFile(file_path)

while linePosition < len(arrayLines):
  if titleToSearch in arrayLines[linePosition]:
    if pageOrPosition == "1":
      position = getPosition(arrayLines, linePosition)
    else:
      position = getPage(arrayLines, linePosition)
    notes = ifIsNote(arrayLines[linePosition+1], arrayLines[linePosition+3])
    dictPositionNotes.update(ordernateNotesNHighlights(position, notes))          
  linePosition+=1

document = initDocument()
document.add_heading(titleToSearch)

otherDict = {}
otherDict = dict(sorted(dictPositionNotes.items(), key=operator.itemgetter(1)))
lastDict = markRepeatedNotes(otherDict)
dictFinal = cleanDictionary(lastDict)

for notes in dictFinal.keys():
    if "== Nota Pessoal: " in notes:
      paragraph = document.add_paragraph()
      paragraph.add_run(notes).bold = True
    else:
      document.add_paragraph(notes)

alphanumeric_filter = filter(str.isalnum, titleToSearch)
cleanTitle = "".join(alphanumeric_filter)

print("Separação entre os dicionários\n \n")
print(dictFinal)

document.save('C:/Users/luiz.sa/Desktop/ExtrairNotasKindle/'+cleanTitle+'.docx')

print("O programa foi um sucesso!")